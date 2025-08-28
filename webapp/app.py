from flask import send_from_directory
def enrich_plan_data(plan_data: dict) -> dict:
    """Add mission summary and ensure all relevant fields are present in the output JSON."""
    plan_data = dict(plan_data)  # shallow copy
    plan_data["kensho_mission"] = (
        "Kensho bridges the gap between unstructured documents and structured project plans, automating days of manual work into minutes. "
        "Upload any project brief, and Kensho will analyze, parse, and deliver a structured plan—saving you time and effort."
    )
    plan_data["generated_at"] = datetime.now().isoformat()
    return plan_data
# webapp/app.py
import json
import logging
import os
import subprocess
import sys
import threading
import uuid
from datetime import datetime
from typing import Any, Dict

# Document parsing libraries
import PyPDF2
import openpyxl
from docx import Document

# Add the project root to the Python path to allow imports from kensho_engine
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from flask import Flask, jsonify, render_template, request  # noqa: E402

from Kensho_engine.brain import analyze_document_text  # noqa: E402

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(name)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

app = Flask(__name__)
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), "uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16MB max file size

# Store for async task status
task_status: Dict[str, Dict[str, Any]] = {}
task_lock = threading.Lock()

# Allowed file extensions and MIME types for security
ALLOWED_EXTENSIONS = {"txt", "pdf", "docx", "xlsx"}
ALLOWED_MIME_TYPES = {
    "text/plain",
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/msword",  # For older .doc files if uploaded as .docx
}


def allowed_file(filename: str, mimetype: str) -> bool:
    """Check if uploaded file is allowed based on extension and MIME type"""
    if not filename:
        return False

    # Check file extension
    if "." not in filename:
        return False

    extension = filename.rsplit(".", 1)[1].lower()
    if extension not in ALLOWED_EXTENSIONS:
        return False

    # Check MIME type
    if mimetype not in ALLOWED_MIME_TYPES:
        return False

    return True


def extract_text_from_file(file, filename: str) -> str:
    """Extract text content from uploaded file based on file type"""
    file_ext = filename.rsplit(".", 1)[1].lower()

    try:
        if file_ext == "txt":
            return file.read().decode("utf-8")

        elif file_ext == "pdf":
            # Reset file pointer to beginning
            file.seek(0)
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()

        elif file_ext == "docx":
            # Reset file pointer to beginning
            file.seek(0)
            doc = Document(file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()

        elif file_ext == "xlsx":
            # Reset file pointer to beginning
            file.seek(0)
            workbook = openpyxl.load_workbook(file, data_only=True)
            text = ""

            # Extract text from all worksheets
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"\n=== {sheet_name} ===\n"

                for row in sheet.iter_rows(values_only=True):
                    row_text = []
                    for cell in row:
                        if cell is not None:
                            row_text.append(str(cell))
                    if row_text:  # Only add non-empty rows
                        text += " | ".join(row_text) + "\n"

            return text.strip()

        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

    except Exception as e:
        logger.error(f"Error extracting text from {filename}: {e}")
        raise ValueError(f"Failed to extract text from {filename}: {str(e)}")


def validate_file_content(content: str) -> bool:
    """Validate file content for security"""
    if not content or not content.strip():
        return False

    # Check file size (additional check)
    if len(content.encode("utf-8")) > 1024 * 1024:  # 1MB max content
        return False

    # Basic content validation - ensure it's readable text
    try:
        content.encode("utf-8")
        return True
    except UnicodeError:
        return False


@app.route("/")
def index():
    """Renders the main page."""
    return render_template("index.html")


@app.route("/analyze", methods=["POST"])
def analyze():
    """
    Handles file upload, calls the Brain to analyze it,
    and returns the structured JSON with enhanced security.
    """
    logger.info("Received analysis request")

    if "document" not in request.files:
        logger.warning("No file part in request")
        return jsonify({"error": "No file part"}), 400

    file = request.files["document"]
    if file.filename == "":
        logger.warning("No file selected")
        return jsonify({"error": "No selected file"}), 400

    # Security validation
    if not allowed_file(file.filename, file.mimetype):
        logger.warning(f"Invalid file type: {file.filename}, MIME: {file.mimetype}")
        return jsonify({"error": "Only .txt, .pdf, .docx, and .xlsx files are allowed"}), 400

    try:
        # Extract text based on file type
        content = extract_text_from_file(file, file.filename)

        # Validate content
        if not validate_file_content(content):
            logger.warning("Invalid file content")
            return jsonify({"error": "Invalid file content"}), 400

        project_title = os.path.splitext(file.filename)[0].replace("_", " ").title()

        logger.info(f"Analyzing document: {project_title}")

        # Call the real Brain logic with enhanced error handling
        plan_data = analyze_document_text(content, project_title)
        plan_data = enrich_plan_data(plan_data)

        logger.info("Document analysis completed successfully")
        return jsonify(plan_data)
    except UnicodeDecodeError:
        logger.error("File encoding error")
        return jsonify({"error": "File must be UTF-8 encoded text"}), 400
    except ValueError as e:
        logger.error(f"Validation error: {e}")
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        logger.error(f"Analysis error: {e}")
        return jsonify({"error": f"Analysis failed: {str(e)}"}), 500

# --- Save as Word/Text endpoint ---
@app.route("/save_local", methods=["POST"])
def save_local():
    """Save the analyzed plan as a Word (.docx) or text file and provide a download link."""
    try:
        data = request.json
        plan = data.get("plan")
        if not plan:
            return jsonify({"error": "No plan data provided"}), 400

        plan = enrich_plan_data(plan)
        project_name = plan.get("project_name", "Kensho_Plan").replace(" ", "_")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"{project_name}_{timestamp}"

        # Save as .docx with enhanced format
        docx_path = os.path.join(app.config["UPLOAD_FOLDER"], base_filename + ".docx")
        doc = Document()
        doc.add_heading(plan.get("projectName", plan.get("project_name", "Kensho Project")), 0)
        
        # Add project metadata if available
        if plan.get("projectObjective"):
            doc.add_heading("Project Objective", level=1)
            doc.add_paragraph(plan["projectObjective"])
        
        # Add team information if available
        if plan.get("team"):
            doc.add_heading("Project Team", level=1)
            for member in plan["team"]:
                doc.add_paragraph(f"• {member.get('memberName', '')}: {member.get('role', '')} ({member.get('level', '')})", style="List Bullet")
        
        # Add phases if available
        if plan.get("phases"):
            doc.add_heading("Project Phases", level=1)
            for phase in plan["phases"]:
                doc.add_heading(f"{phase.get('phaseName', '')}", level=2)
                doc.add_paragraph(f"Owner: {phase.get('phaseOwner', 'Not Specified')}")
                doc.add_paragraph(f"Status: {phase.get('phaseStatus', 'Pending')}")
                if phase.get("subTasks"):
                    doc.add_paragraph("Sub-tasks:")
                    for subtask in phase["subTasks"]:
                        doc.add_paragraph(f"• {subtask.get('taskName', '')}", style="List Bullet")
        
        # Add requirements if available
        if plan.get("requirements"):
            doc.add_heading("Requirements", level=1)
            for req in plan["requirements"]:
                doc.add_paragraph(f"{req.get('reqId', '')}: {req.get('description', '')}", style="List Bullet")
        
        # Add thematic groups (for backward compatibility)
        doc.add_heading("Detailed Analysis", level=1)
        doc.add_paragraph(plan.get("kensho_mission", ""))
        doc.add_paragraph(f"Generated at: {plan.get('generated_at', '')}")
        doc.add_paragraph("")
        for group in plan.get("thematic_groups", []):
            doc.add_heading(group.get("group_name", "Unnamed Group"), level=2)
            if group.get("group_description"):
                doc.add_paragraph(group["group_description"])
            for task in group.get("tasks", []):
                doc.add_paragraph(f"- {task.get('task_name', '')}", style="List Bullet")
                if task.get("owner"):
                    doc.add_paragraph(f"  Owner: {task['owner']}", style="Intense Quote")
                if task.get("details"):
                    doc.add_paragraph(f"  Details: {task['details']}", style="Intense Quote")
        doc.save(docx_path)

        # Save as .txt with enhanced professional format
        txt_path = os.path.join(app.config["UPLOAD_FOLDER"], base_filename + ".txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            # Use professional analysis if available, otherwise fall back to basic format
            if plan.get("professional_analysis"):
                f.write(plan["professional_analysis"])
            else:
                # Fallback to basic format
                f.write(plan.get("projectName", plan.get("project_name", "Kensho Project")) + "\n")
                f.write(plan.get("kensho_mission", "") + "\n")
                f.write(f"Generated at: {plan.get('generated_at', '')}\n\n")
                for group in plan.get("thematic_groups", []):
                    f.write(f"# {group.get('group_name', 'Unnamed Group')}\n")
                    if group.get("group_description"):
                        f.write(group["group_description"] + "\n")
                    for task in group.get("tasks", []):
                        f.write(f"- {task.get('task_name', '')}\n")
                        if task.get("owner"):
                            f.write(f"  Owner: {task['owner']}\n")
                        if task.get("details"):
                            f.write(f"  Details: {task['details']}\n")
                    f.write("\n")

        # Default to docx for download
        download_url = f"/download/{os.path.basename(docx_path)}"
        return jsonify({"download_url": download_url})
    except Exception as e:
        logger.error(f"Error saving local file: {e}")
        return jsonify({"error": str(e)}), 500

# --- Serve generated files ---
@app.route("/download/<filename>")
def download_file(filename):
    """Serve a file from the uploads folder."""
    return send_from_directory(app.config["UPLOAD_FOLDER"], filename, as_attachment=True)


def execute_hands_async(task_id: str, plan_data: Dict[str, Any], target: str, json_path: str):
    """Execute hands orchestrator asynchronously"""
    with task_lock:
        task_status[task_id]["status"] = "running"
        task_status[task_id]["message"] = "Execution in progress..."

    try:
        # Get the project root directory
        project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))

        # Construct the command to run the hands orchestrator as a module
        # Use absolute paths and validate input
        command = [
            sys.executable,
            "-m",
            "Kensho_engine.hands",
            "--input",
            os.path.abspath(json_path),
            "--target",
            target,
            "--config",
            os.path.abspath(os.path.join(project_root, "config.ini")),
        ]

        logger.info(f"Executing command for task {task_id}: {' '.join(command)}")

        result = subprocess.run(
            command,
            capture_output=True,
            text=True,
            check=True,
            cwd=project_root,  # Run from the project root
            timeout=300,  # 5 minute timeout
        )

        with task_lock:
            task_status[task_id]["status"] = "completed"
            task_status[task_id]["success"] = True
            task_status[task_id]["message"] = f"Successfully executed for target: {target}"
            task_status[task_id]["logs"] = result.stdout
            task_status[task_id]["completed_at"] = datetime.now().isoformat()

        logger.info(f"Task {task_id} completed successfully")

    except subprocess.TimeoutExpired:
        with task_lock:
            task_status[task_id]["status"] = "failed"
            task_status[task_id]["success"] = False
            task_status[task_id]["message"] = f"Execution timed out for target: {target}"
            task_status[task_id]["completed_at"] = datetime.now().isoformat()
        logger.error(f"Task {task_id} timed out")

    except subprocess.CalledProcessError as e:
        with task_lock:
            task_status[task_id]["status"] = "failed"
            task_status[task_id]["success"] = False
            task_status[task_id]["message"] = f"Execution failed for target: {target}"
            task_status[task_id]["logs"] = f"STDOUT:\n{e.stdout}\n\nSTDERR:\n{e.stderr}"
            task_status[task_id]["completed_at"] = datetime.now().isoformat()
        logger.error(f"Task {task_id} failed with exit code {e.returncode}")

    except Exception as e:
        with task_lock:
            task_status[task_id]["status"] = "failed"
            task_status[task_id]["success"] = False
            task_status[task_id]["message"] = f"Unexpected error: {str(e)}"
            task_status[task_id]["completed_at"] = datetime.now().isoformat()
        logger.error(f"Task {task_id} failed with unexpected error: {e}")


@app.route("/execute", methods=["POST"])
def execute():
    """
    Receives the plan JSON and a target platform,
    then calls the Hands orchestrator asynchronously.
    """
    logger.info("Received execution request")

    try:
        data = request.json
        if not data:
            return jsonify({"error": "No JSON data provided"}), 400

        plan_data = data.get("plan")
        target = data.get("target")

        if not plan_data or not target:
            logger.warning("Missing plan data or target")
            return jsonify({"error": "Missing plan data or target"}), 400

        # Validate target
        valid_targets = ["jira", "asana", "confluence", "trello", "slack"]
        if target not in valid_targets:
            logger.warning(f"Invalid target: {target}")
            return jsonify({"error": f"Invalid target. Must be one of: {valid_targets}"}), 400

        # Generate unique task ID
        task_id = str(uuid.uuid4())

        # Save plan data to temporary file with unique name
        json_filename = f"temp_plan_{task_id}.json"
        json_path = os.path.join(app.config["UPLOAD_FOLDER"], json_filename)

        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(plan_data, f, ensure_ascii=False, indent=4)

        # Initialize task status
        with task_lock:
            task_status[task_id] = {
                "status": "pending",
                "target": target,
                "created_at": datetime.now().isoformat(),
                "message": "Task queued for execution",
            }

        # Start async execution
        thread = threading.Thread(target=execute_hands_async, args=(task_id, plan_data, target, json_path))
        thread.daemon = True
        thread.start()

        logger.info(f"Started async execution for task {task_id}, target: {target}")

        return jsonify({"task_id": task_id, "status": "pending", "message": "Task queued for execution"})

    except Exception as e:
        logger.error(f"Error in execute endpoint: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/status/<task_id>", methods=["GET"])
def get_task_status(task_id: str):
    """Get the status of an async task"""
    with task_lock:
        if task_id not in task_status:
            return jsonify({"error": "Task not found"}), 404

        status = task_status[task_id].copy()

    return jsonify(status)


if __name__ == "__main__":
    # Production configuration
    port = int(os.environ.get("PORT", 5001))
    debug = os.environ.get("FLASK_ENV") != "production"
    host = "0.0.0.0" if not debug else "127.0.0.1"

    logger.info(f"Starting Flask app on {host}:{port} (debug={debug})")
    app.run(host=host, port=port, debug=debug)
